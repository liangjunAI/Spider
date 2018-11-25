#-*- coding:utf-8 -*-
'''
    spider project: 政府网站有关人社信息数据爬取
    url: http://www.gov.cn/zhengce/index.htm
    date： 2017/10/17
    author:huliangjun
'''
import urllib2
from lxml import etree
import sys
import copy
import datetime
import os
import urllib
from docx.shared import Inches
import requests as req
import cStringIO
import time as t
import pandas as pd
from docx import Document
import pymongo

def get_content(url):
    headers = {"User-Agent":"Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/45.0.2454.101 Safari/537.36"}
    try:
        req = urllib2.Request(url,headers=headers)
        content = urllib2.urlopen(req,timeout=1000).read()
        content = etree.HTML(content)
        return content
    except (urllib2.URLError,Exception),e:
        if hasattr(e,'reason'):
            print '抓取失败，具体原因：',e.reason
            for i in range(5):
                t.sleep(5)
                rep = urllib2.urlopen(req,timeout=1000)
                content = rep.read().decode('utf-8')
            return content

def get_datetime():
    times = datetime.datetime.now()
    year = times.year
    month = times.month
    day = times.day
    return  year,month,day

def get_info(collection,keyword,num_id,content,year,month,day):
    if not os.path.isdir("F:\\GOV"):
        os.mkdir("F:\\GOV")
    num_id = int(num_id)
    num_id = num_id + 1
    tom = str(num_id)
    leng = tom.__len__()
    if leng < 6:
        num_id = "0"*(6-leng) + str(num_id)
    tom = "GOV_CN" + "-(" + str(year) + str(month) + str(day) + ")-" + str(num_id)
    filename = "F:\\GOV\\" + tom + ".docx"
    info_body = content.xpath('.//div[@class="gov-right"]/div[@class="result"]//li[@class="res-list"]')
    for i in range(len(info_body)):
        html = info_body[i].xpath('./h3//@href')[0]
        time = info_body[i].xpath('./p/span[@class="sp sourceName"]/text()')[1]
        print 'html:',html
        content_total = ""
        cursor = collection.find({"_id": html})
        # if html not in url_list:
        if cursor.count() == 0:
            # url_list.append(html)
            # print 'url_list:',url_list
            title = info_body[i].xpath('./h3//text()')[0]
            print 'title:',title
            test_body = get_content(html)
            content_info = get_article(content_total,test_body,filename=filename)
            info_db = info_dict(html,title,time,keyword,content_info,tom)
            info_csv = info_save(url,title,time,keyword,tom)
            info_list.append(info_csv)
            collection.save(dict(info_db))
            print 'info:',info_db
            print 'info_csv:',info_csv
            print 'info_list:', info_list

    return content_total,num_id,html

def info_save(url,title,time,tabname,tom):
    save = {}
    save["num_id"] = tom
    save["title"] = title
    save["_id"] = url
    save["time"] = time
    save["Source"] = "中华人民共和国中央人民政府"
    save["tab"] = tabname
    return save

def info_dict(url,title,time,tabname,content,tom):
    Info = {}
    Info["_id"] = url
    Info["Source"] = "中华人民共和国中央人民政府"
    Info["title"] = title
    Info["time"] = time
    Info["tab"] = tabname
    Info["content"] = content.lstrip()
    Info["num_id"] = tom
    return Info

def get_article(content_total,url_body,filename):
    text = url_body.xpath('.//td[@class="b12c"]/p|.//td[@class="b12c"]/div|.//td[@class="b12c"]/span|.//td[@class="b12c"]/font')
    doc = Document()
    for x in text:
        row = x.xpath('.//tbody//tr')
        try:
            if row:
                for col in row:
                    table = doc.add_table(rows=1, cols=int(len(col)), style='Table Grid')
                    hdr_cells = table.rows[0].cells
                    td = col.xpath('./td')
                    td_list = []
                    for t in td:
                        ins_data = t.xpath('.//text()')
                        ins_data = filter(lambda x: x != '\r\n      ', ins_data)
                        tmp = ""
                        for i in range(len(ins_data)):
                            tmp += ins_data[i]
                        td_list.append(tmp)
                    length = len(td_list)
                    for i in range(length):
                        hdr_cells[i].text = td_list[i]
            else:
                img_ads = x.xpath('.//@src')
                imglist = ""
                if img_ads:
                    for h in img_ads:
                        if h.startswith('http://'):
                            try:
                                img = cStringIO.StringIO(urllib2.urlopen(h).read())
                            except:
                                img = cStringIO.StringIO(req.get(h).content)
                        else:
                            ind = html.rfind('/') + 1
                            h = html.replace(html[ind:], h)
                            try:
                                img = cStringIO.StringIO(urllib2.urlopen(h).read())
                            except:
                                img = cStringIO.StringIO(req.get(h).content)
                        doc.add_picture(img, width=Inches(4.25))
                        imglist = imglist + h + " "
                paragraph = x.xpath('.//text()')
                str_ = ""
                for j in paragraph:
                    str_ = str_ + j
                doc.add_paragraph(str_)
                content_total = content_total + imglist + str_ + "\n"
        except:
            pass
    doc.save(filename)
    return content_total

def get_page(content):
    page = content.xpath('.//div[@class="content"]//div[@id="page" and @class = "page"]/text()')
    page = int(page[-2].strip().strip(u'\xa0').strip(u'\u5171').strip(u'\u9875'))
    return page

if __name__ == '__main__':
    client = pymongo.MongoClient("192.168.46.231", 27017)
    collection = client["spider"]["zhongyangzhengfu"]
    content_words = ['人社','生育','失业','工伤','社保','医保','养老']
    title_words = ['生育','失业','工伤','养老']
    num_id = 0
    year,month,day = get_datetime()
    url_list = []
    info_list = []
    input_csv = "F:\\GOV\\" + "GOV_CN" + "_(" + str(year) + str(month) + str(day) + ")" + ".csv"
    for i in content_words:
        url = 'http://sousuo.gov.cn/s.htm?t=paper&advance=true&title=&content='+urllib.quote(i)
        t.sleep(1)
        content = get_content(url)
        page = get_page(content)
        print url
        print page
        if page > 1:
            for p in range(page):
                url_tmp = 'http://sousuo.gov.cn/s.htm?q=&n=10&p={0}&t=paper&advance=true&title=&content='+urllib.quote(i)
                url_content = url_tmp.format(p)
                content_page = get_content(url_content)
                content_, num_id, html = get_info(collection,i,num_id,content_page,year,month,day)
        else:
            content_, num_id, html = get_info(collection,i,num_id,content,year,month,day)

    for i in title_words:
        url = 'http://sousuo.gov.cn/s.htm?t=paper&advance=true&title='+urllib.quote(i)
        content = get_content(url)
        page = get_page(content)
        if page > 1:
            for p in range(page):
                url_tmp = 'http://sousuo.gov.cn/s.htm?q=&n=10&p={0}&t=paper&advance=true&title=' + urllib.quote(i)
                url_title = url_tmp.format(p)
                t.sleep(1)
                content_page = get_content(url_title)
                content_, num_id, html = get_info(collection,i,num_id,content_page,year,month,day)
        else:
            content_, num_id, html = get_info(collection,i,num_id,content,year,month,day)

    info_csv = pd.DataFrame(info_list)
    info_csv.to_csv(input_csv,index=None,header=None,encoding='utf-8')


