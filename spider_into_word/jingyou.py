#-*- encoding:utf-8 -*-

import urllib2
from lxml import etree
import docx
import sys
import copy
import requests as req
import datetime
import os
import urllib
import cStringIO
import time as t
import pymongo
print "https://www.joyowo.com/zixun/"

def get_Soup(url):
    req = urllib2.Request(url)
    req.add_header("User-Agent","Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/45.0.2454.101 Safari/537.36")
    content = urllib2.urlopen(req,timeout=1000).read()
    # print BeautifulSoup(content, 'lxml')
    content = etree.HTML(content)
    return content

def get_content(num_id,source_url,tabname,collection,year,month,day):
    break_num = 1
    break_flag = 1
    # start_url = source_url % 1
    # content = get_Soup(start_url)
    # info_body = content.xpath('//li[@class="page-total"]//text()')[0]
    # total = int(info_body.split("/")[1])
    total = 10
    for j in range(1,total):
        if break_flag == 1:
            t.sleep(0.5)
            start_url = source_url%j
            content = get_Soup(start_url)
            info_body = content.xpath('//div[@class="article-info"]')
            for i in info_body:
                print tabname + "   " + str(num_id)
                num_id = num_id + 1
                try:
                    url = "https://www.joyowo.com" + i.xpath('div[@class="article-title"]//@href')[0]
                    cursor = collection.find({"_id": url})
                    if cursor.count()!= 0:
                        break_num += 1
                        if break_num > 3:
                            # client.close()
                            break_flag = 0
                            break
                        next()
                    title = i.xpath('div[@class="article-title"]//text()')[0]
                    abstracts = i.xpath('p[@class="article-content"]//text()')
                    if len(abstracts) > 0:
                        abstracts = abstracts[0]
                    else:
                        abstracts = ""
                    time = i.xpath('div[@class="article-publish"]//text()')[0]
                    Info={}
                    Info["_id"] = url
                    Info["Source"]="金柚网"
                    Info["tab"]= tabname
                    Info["time"] = time
                    Info["title"] = title
                    Info["abstracts"] = abstracts
                    end,tom =get_aim(num_id,url,year,month,day,title)
                    Info["content"] = end.lstrip()
                    Info["num_id"] = tom
                    collection.save(dict(Info))
                    # filename = "D:\\jinjuwang\\" + str(num_id) + ".docx"
                    # file = docx.Document()
                    # file.add_paragraph(end)
                    # file.save(filename)
                except :
                    pass
    return num_id
def get_datetime():
    times = datetime.datetime.now()
    year = times.year
    month = times.month
    day = times.day
    return  year,month,day


def get_aim(num_id,url,year,month,day,title):
    if not os.path.isdir("D:\\jinjuwang"):
        os.mkdir("D:\\jinjuwang")
    tom = str(num_id)
    leng = tom.__len__()
    if leng < 6:
        num_id = "0"*(6-leng) + str(num_id)
    tom = "JYW" + "-(" + str(year) + str(month) + str(day) + ")-" + str(num_id)
    filename = "D:\\jinjuwang\\" + tom + ".docx"
    file = docx.Document()
    file.add_paragraph(title)
    aim_content = get_Soup(url)
    content = ""
    body = aim_content.xpath('//div[@class="article-info"]//p|//div[@class="article-info"]//center')
    print len(body)
    for i in body:
        try:
            img_url = i.xpath('.//@src')
            imglist = ""
            if len(img_url) > 0:
                for h in img_url:
                    try:
                        img = cStringIO.StringIO(urllib2.urlopen(h).read())
                    except:
                        img = cStringIO.StringIO(req.get(h).content)
                    #在文档里面添加图片
                    file.add_picture(img)
                    imglist = imglist + h + " "
            text = i.xpath(".//text()")
            string_ = ""
            for j in text:
                string_ = string_ + j
            file.add_paragraph(string_)
            content = content + string_ + imglist + "\n"
        except:
            pass
    file.save(filename)
    return content,tom



if __name__ == "__main__":
    client = pymongo.MongoClient("192.168.46.231", 27017)
    collection = client["spider"]["jinyouwang2"]
    num_id = 0
    urllist = [
        ["https://www.joyowo.com/shebaozhengce_%d/", "社保政策"],
        ["https://www.joyowo.com/shebaonews_%d/", "社保新闻"],
        ["https://www.joyowo.com/shebaobanli_%d/", "社保办理"],
        ["https://www.joyowo.com/shebaozhishi_%d/", "社保知识"],
        ["https://www.joyowo.com/shebaoka_%d/", "社保卡"],
        ["https://www.joyowo.com/yanglaobaoxian_%d/", "养老保险"],
        ["https://www.joyowo.com/yiliaobaoxian_%d/", "医疗保险"],
        ["https://www.joyowo.com/shengyubaoxian_%d/", "生育保险"],
        ["https://www.joyowo.com/gongshangbaoxian_%d/", "工伤保险"],
        ["https://www.joyowo.com/shiyebaoxian_%d/", "失业保险"],
        ["https://www.joyowo.com/gongjijin_%d/", "住房公积金"],
        ["https://www.joyowo.com/shebaoanli_%d/", "社保案例"]
    ]
    year, month, day = get_datetime()
    for i in urllist:
        num_id=get_content(num_id,i[0],i[1],collection,year,month,day)
    client.close()



