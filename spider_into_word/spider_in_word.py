#-*- coding:utf-8 -*-
'''
    date： 2017/10/17
    author:liangjun
'''
import urllib2
from lxml import etree
from docx.shared import Inches
import requests as req
import cStringIO
import time as t
from docx import Document

def get_content(url):
    headers = {"User-Agent":"Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/45.0.2454.101 Safari/537.36"}
    try:
        req = urllib2.Request(url,headers=headers)
        content = urllib2.urlopen(req,timeout=1000).read()
        content = etree.HTML(content)
        return content
    except (urllib2.URLError,Exception),e:
        if hasattr(e,'reason'):
            print '抓取失败，具体原因：',e.reason
            for i in range(5):
                t.sleep(5)
                rep = urllib2.urlopen(req,timeout=1000)
                content = rep.read().decode('utf-8')
            return content

def get_article(url_body,filename,html):
    content_total = ""
    text = url_body.xpath('.//td[@class="b12c"]/p|.//td[@class="b12c"]/div|.//td[@class="b12c"]/span|.//td[@class="b12c"]/font')
    doc = Document()
    for x in text:
        row = x.xpath('.//tbody//tr')
        try:
            if row:
                for col in row:
                    table = doc.add_table(rows=1, cols=int(len(col)), style='Table Grid')
                    hdr_cells = table.rows[0].cells
                    td = col.xpath('./td')
                    td_list = []
                    for t in td:
                        ins_data = t.xpath('.//text()')
                        ins_data = filter(lambda x: x != '\r\n      ', ins_data)
                        tmp = ""
                        for i in range(len(ins_data)):
                            tmp += ins_data[i]
                        td_list.append(tmp)
                    length = len(td_list)
                    for i in range(length):
                        hdr_cells[i].text = td_list[i]
            else:
                img_ads = x.xpath('.//@src')
                imglist = ""
                if img_ads:
                    for h in img_ads:
                        if h.startswith('http://'):
                            try:
                                img = cStringIO.StringIO(urllib2.urlopen(h).read())
                            except:
                                img = cStringIO.StringIO(req.get(h).content)
                        else:
                            ind = html.rfind('/') + 1
                            h = html.replace(html[ind:], h)
                            try:
                                img = cStringIO.StringIO(urllib2.urlopen(h).read())
                            except:
                                img = cStringIO.StringIO(req.get(h).content)
                        doc.add_picture(img, width=Inches(4.25))
                        imglist = imglist + h + " "
                paragraph = x.xpath('.//text()')
                str_ = ""
                for j in paragraph:
                    str_ = str_ + j
                doc.add_paragraph(str_)
                content_total = content_total + imglist + str_ + "\n"
        except:
            pass
    doc.save(filename)
    return content_total

def get_main():
    url = 'http://www.gov.cn/zhengce/content/2017-03/01/content_5172013.htm'
    filePath = 'spider_test.docx'

    content = get_content(url)
    paper_cotent = get_article(content, filePath,url)

if __name__ == '__main__':
    get_main()

